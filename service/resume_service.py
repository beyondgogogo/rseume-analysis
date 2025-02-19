import logging
import os
from io import BytesIO

import PyPDF2
import re
from datetime import datetime
from docx import Document  # 处理 Word 文件
import easyocr
from PIL import Image  # 处理图片

from utils import glm  # 假设你使用的 GLM 模型

# 设置日志
logging.basicConfig(level=logging.DEBUG)

class ResumeService:

    # 判断文件类型
    def detect_file_type(self, file):
        # 获取文件的扩展名（小写）
        file_extension = os.path.splitext(file.filename)[1].lower()

        if file_extension == '.pdf':
            return 'pdf'
        elif file_extension in ['.doc', '.docx']:
            return 'word'
        elif file_extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']:
            return 'image'
        else:
            raise ValueError(f"不支持的文件类型：{file_extension}")

    # 提取PDF文件中的文本
    def extract_text_from_pdf(self, file):
        try:
            logging.info(f"开始提取PDF文本：{file.name if hasattr(file, 'name') else '未命名文件'}")
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                if page_text:
                    text += page_text
                    logging.debug(f"成功提取第 {page_num + 1} 页的文本")
                else:
                    logging.warning(f"第 {page_num + 1} 页未提取到文本")
            if not text:
                logging.error("提取的文本为空，请检查PDF文件是否有效或格式正确。")
                raise ValueError("提取的PDF文本为空")
            logging.info("成功提取PDF文本")
            return text
        except Exception as e:
            logging.error(f"提取PDF文本时发生错误：{str(e)}")
            raise Exception(f"提取PDF文本时发生错误：{str(e)}")

    # 提取Word文件中的文本
    def extract_text_from_word(self, file):
        try:
            logging.info(f"开始提取Word文本：{file.name if hasattr(file, 'name') else '未命名文件'}")
            # 确保文件对象被处理正确
            file_content = file.read()
            file_stream = BytesIO(file_content)
            doc = Document(file_stream)
            text = ""

            # 提取段落中的文本
            for para in doc.paragraphs:
                if para.text:
                    text += para.text + '\n'

            # 提取表格中的文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            text += cell_text + '\n'

            logging.info("成功提取Word文本")
            logging.info("text的内容" + text)
            if not text.strip():
                logging.warning("提取的文本为空，请检查Word文件内容。")
            return text
        except Exception as e:
            logging.error(f"提取Word文本时发生错误：{str(e)}")
            raise Exception(f"提取Word文本时发生错误：{str(e)}")

    # 提取图片中的文本
    def extract_text_from_image(self, file):
        try:
            # 创建 OCR 识别器
            reader = easyocr.Reader(['en', 'ch_sim'])  # 支持英文和简体中文

            # 创建一个临时文件夹保存图片
            temp_folder = "temp_images"
            if not os.path.exists(temp_folder):
                os.makedirs(temp_folder)

            # 保存图片到本地
            file_path = os.path.join(temp_folder, file.filename)
            file.save(file_path)
            logging.info(f"图片已保存至：{file_path}")

            # 使用PIL加载图像
            image = Image.open(file_path)

            # 在图像对象上进行 OCR 识别
            result = reader.readtext(image)

            # 输出识别的文本
            extracted_text = ""
            for detection in result:
                extracted_text += detection[1] + "\n"

            # 删除临时文件
            os.remove(file_path)
            logging.info(f"临时文件已删除：{file_path}")

            return extracted_text
        except Exception as e:
            logging.error(f"提取图片文本时发生错误：{str(e)}")
            raise Exception(f"提取图片文本时发生错误：{str(e)}")

    # 调用GLM模型进行简历分块
    def split_resume_with_glm(self, extracted_text):
        try:
            logging.info("开始调用GLM模型分块简历内容")
            prompt = f"请将以下简历内容分成几个部分，包括但不限于：个人基本信息、教育经历、实习经历、工作经历、项目经历、相关技能等。\n\n{extracted_text}"

            # 调用GLM模型并返回分块结果
            result = glm.chat_with_ai("user", prompt)
            logging.info("简历内容分块成功")
            # 进行格式检查
            if not result or not isinstance(result, str):
                logging.error("返回的简历分块内容无效或格式不正确")
                raise ValueError("GLM返回的简历分块内容无效")
            return result
        except Exception as e:
            logging.error(f"调用GLM模型分块简历内容时发生错误：{str(e)}")
            raise Exception(f"调用GLM模型分块简历内容时发生错误：{str(e)}")

    # 清理简历内容中的Markdown标记
    def clean_markdown(self, content):
        content = re.sub(r'###\s+', '', content)
        content = re.sub(r'-\s+', '', content)
        return content.strip()

    # 保存每个简历部分为单独的txt文件
    def save_resume_to_txt(self, resume_id, resume_parts):
        try:
            parent_folder = "resume_list"
            if not os.path.exists(parent_folder):
                os.makedirs(parent_folder)
                logging.info(f"创建文件夹：{parent_folder}")

            folder_name = os.path.join(parent_folder, f"resume_{resume_id}")
            if not os.path.exists(folder_name):
                os.makedirs(folder_name)
                logging.info(f"创建文件夹：{folder_name}")
            else:
                logging.warning(f"文件夹 {folder_name} 已存在")

            resume_data = resume_parts.split("\n")
            current_section = ""
            current_content = []

            section_titles = ["个人基本信息", "教育经历", "实习经历", "项目经历", "相关技能"]

            for part in resume_data:
                part = part.strip()
                for title in section_titles:
                    if part.startswith(f"### {title}"):  # 匹配到标题
                        if current_section:
                            cleaned_content = self.clean_markdown("\n".join(current_content))
                            file_name = os.path.join(folder_name, f"{current_section}.txt")
                            with open(file_name, "w", encoding="utf-8") as f:
                                f.write(cleaned_content)
                            logging.debug(f"已保存文件：{file_name}")

                        current_section = title
                        current_content = [part]
                        logging.debug(f"检测到新部分标题：{current_section}")
                        break
                else:
                    if part:
                        current_content.append(part)

            if current_section:
                cleaned_content = self.clean_markdown("\n".join(current_content))
                file_name = os.path.join(folder_name, f"{current_section}.txt")
                with open(file_name, "w", encoding="utf-8") as f:
                    f.write(cleaned_content)
                logging.debug(f"已保存文件：{file_name}")

            logging.info(f"简历内容成功保存到文件夹：{folder_name}")
            return folder_name
        except Exception as e:
            logging.error(f"保存简历到txt文件时发生错误：{str(e)}")
            raise Exception(f"保存简历到txt文件时发生错误：{str(e)}")

    # 处理简历的主流程
    def process_resume(self, file):
        try:
            logging.info("开始处理简历流程")

            # 判断文件类型
            file_type = self.detect_file_type(file)
            logging.info(f"文件类型：{file_type}")

            # 根据文件类型提取文本
            if file_type == 'pdf':
                extracted_text = self.extract_text_from_pdf(file)
            elif file_type == 'word':
                extracted_text = self.extract_text_from_word(file)
                print("111")
                print(extracted_text)
            elif file_type == 'image':
                extracted_text = self.extract_text_from_image(file)
            else:
                raise ValueError("无法识别的文件类型")

            # 调用GLM模型进行简历分块
            resume_parts = self.split_resume_with_glm(extracted_text)

            # 生成当前时间戳作为简历号
            resume_id = datetime.now().strftime("%Y%m%d%H%M%S")
            logging.info(f"生成的简历ID为：{resume_id}")

            # 保存简历内容到txt文件夹中
            folder_name = self.save_resume_to_txt(resume_id, resume_parts)

            logging.info("简历处理成功")
            return {"message": f"简历上传并处理成功，简历号：{resume_id}", "folder": folder_name}
        except Exception as e:
            logging.error(f"处理简历时发生错误：{str(e)}")
            return {"message": f"处理简历时发生错误：{str(e)}"}
